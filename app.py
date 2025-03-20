import streamlit as st
import markdown
from docx import Document
from io import BytesIO
import re
import base64

def convert_markdown_to_docx(md_text):
    """
    Convert markdown text to a docx file
    
    Parameters:
    md_text (str): Markdown text to convert
    
    Returns:
    BytesIO: A byte stream containing the docx file
    """
    # Convert markdown to HTML
    html = markdown.markdown(md_text, extensions=['tables', 'fenced_code'])
    
    # Create a new Document
    doc = Document()
    
    # Split HTML by various headers and paragraphs
    parts = re.split(r'<h(\d)>(.*?)</h\1>|<p>(.*?)</p>|<pre><code>(.*?)</code></pre>', html, flags=re.DOTALL)
    
    i = 0
    while i < len(parts):
        if parts[i] is None:
            i += 1
            continue
            
        # Header
        if parts[i].isdigit() and i + 1 < len(parts) and parts[i+1] is not None:
            level = int(parts[i])
            text = parts[i+1].strip()
            
            if level == 1:
                doc.add_heading(text, 0)  # Title
            else:
                doc.add_heading(text, level)
            i += 2
            
        # Paragraph
        elif i + 2 < len(parts) and parts[i+2] is not None:
            text = parts[i+2].strip()
            doc.add_paragraph(text)
            i += 3
            
        # Code block
        elif i + 3 < len(parts) and parts[i+3] is not None:
            code = parts[i+3].strip()
            p = doc.add_paragraph()
            run = p.add_run(code)
            run.font.name = 'Courier New'
            i += 4
            
        else:
            # Handle plain text or other elements
            if parts[i].strip():
                doc.add_paragraph(parts[i].strip())
            i += 1

    # Save the document to a BytesIO object
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes

def create_download_link(docx_bytes, filename="document.docx"):
    """
    Create a download link for the docx file
    
    Parameters:
    docx_bytes (BytesIO): Byte stream containing the docx file
    filename (str): Name of the file to download
    
    Returns:
    str: HTML link to download the file
    """
    b64 = base64.b64encode(docx_bytes.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download {filename}</a>'

def main():
    st.title("Markdown to DOCX Converter")
    
    # Example markdown content
    example_md = """# Sample Title
## Heading 2
This is a paragraph with **bold** and *italic* text.

### Code Example
```python
def hello_world():
    print("Hello, World!")
```

## Another Section
- Item 1
- Item 2
- Item 3
"""
    
    # Create a text area for markdown input
    md_text = st.text_area("Enter Markdown", example_md, height=300)
    
    # Add a button to convert markdown to docx
    if st.button("Convert to DOCX"):
        docx_bytes = convert_markdown_to_docx(md_text)
        
        # Display a download link
        st.markdown(create_download_link(docx_bytes), unsafe_allow_html=True)
        
        st.success("Conversion complete! Click the link above to download.")

if __name__ == "__main__":
    main()
