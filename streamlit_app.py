import streamlit as st
import markdown
from bs4 import BeautifulSoup
import base64
import io
import pdfkit
import weasyprint
from docx import Document
from docx.shared import RGBColor
import re

def read_markdown_file(uploaded_file):
    """Read markdown file content from uploaded file"""
    content = uploaded_file.getvalue().decode("utf-8")
    return content

def md_to_html_with_style(md_content):
    """Convert markdown to HTML with styling preserved"""
    # Use extended markdown with extras for better styling
    html = markdown.markdown(
        md_content,
        extensions=[
            'markdown.extensions.extra',
            'markdown.extensions.codehilite',
            'markdown.extensions.tables',
            'markdown.extensions.toc',
            'markdown.extensions.fenced_code'
        ]
    )
    
    # Add some default CSS for styling
    styled_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{ font-family: Arial, sans-serif; line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 20px; }}
            h1 {{ color: #2c3e50; border-bottom: 1px solid #eee; padding-bottom: 10px; }}
            h2, h3, h4 {{ color: #3498db; margin-top: 30px; }}
            pre {{ background-color: #f8f8f8; border: 1px solid #ddd; padding: 10px; border-radius: 5px; overflow-x: auto; }}
            code {{ background-color: #f8f8f8; padding: 2px 5px; border-radius: 3px; font-family: monospace; }}
            blockquote {{ background-color: #f9f9f9; border-left: 4px solid #ccc; padding: 10px 15px; margin: 0; }}
            table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
            th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
            th {{ background-color: #f2f2f2; }}
            tr:nth-child(even) {{ background-color: #f9f9f9; }}
            img {{ max-width: 100%; height: auto; }}
        </style>
    </head>
    <body>
        {html}
    </body>
    </html>
    """
    return styled_html

def convert_to_pdf(html_content):
    """Convert HTML content to PDF using WeasyPrint"""
    try:
        pdf = weasyprint.HTML(string=html_content).write_pdf()
        return pdf
    except Exception as e:
        st.error(f"Error converting to PDF: {e}")
        return None

def convert_to_docx_with_style(html_content):
    """Convert HTML content to DOCX with styling preserved"""
    soup = BeautifulSoup(html_content, 'html.parser')
    doc = Document()
    
    for element in soup.body.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'pre', 'ul', 'ol', 'li', 'blockquote']):
        tag_name = element.name
        
        if tag_name.startswith('h'):
            level = int(tag_name[1])
            heading = doc.add_heading(level=level)
            run = heading.add_run(element.get_text())
            if level == 1:
                run.font.color.rgb = RGBColor(44, 62, 80)  # Dark blue
            else:
                run.font.color.rgb = RGBColor(52, 152, 219)  # Light blue
                
        elif tag_name == 'p':
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(element.get_text())
            
        elif tag_name == 'pre':
            paragraph = doc.add_paragraph()
            code_text = element.get_text()
            run = paragraph.add_run(code_text)
            run.font.name = 'Courier New'
            
        elif tag_name in ['ul', 'ol']:
            # Skip list containers as we'll process individual li elements
            continue
            
        elif tag_name == 'li':
            # Check if this is inside a ul or ol
            parent_tag = element.parent.name
            paragraph = doc.add_paragraph(style='List Bullet' if parent_tag == 'ul' else 'List Number')
            run = paragraph.add_run(element.get_text())
            
        elif tag_name == 'blockquote':
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(element.get_text())
            paragraph.style = 'Quote'
    
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    return docx_io.getvalue()

def create_download_link(content, filename, text):
    """Create a download link for the given content"""
    b64 = base64.b64encode(content).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">{text}</a>'
    return href

def main():
    st.title("Markdown Converter with Style Preservation")
    
    # File upload
    uploaded_file = st.file_uploader("Upload a Markdown file", type=["md", "markdown", "txt"])
    
    if uploaded_file is not None:
        md_content = read_markdown_file(uploaded_file)
        
        # Show preview of uploaded content
        with st.expander("Markdown Content Preview"):
            st.markdown(md_content)
        
        # Output format selection
        output_format = st.selectbox(
            "Select output format",
            ["HTML", "PDF", "DOCX"]
        )
        
        # Style customization option
        with st.expander("Style Customization (HTML only)"):
            custom_css = st.text_area(
                "Custom CSS",
                """body { font-family: Arial, sans-serif; }
h1 { color: #2c3e50; }
h2, h3, h4 { color: #3498db; }""",
                height=200
            )
        
        if st.button("Convert"):
            # First convert to HTML with styling
            html_content = md_to_html_with_style(md_content)
            
            if output_format == "HTML":
                # If custom CSS provided, replace the CSS in the HTML
                if custom_css:
                    html_content = re.sub(
                        r'<style>.*?</style>',
                        f'<style>{custom_css}</style>',
                        html_content,
                        flags=re.DOTALL
                    )
                
                # Provide HTML for download
                st.download_button(
                    label="Download HTML",
                    data=html_content,
                    file_name=f"{uploaded_file.name.split('.')[0]}.html",
                    mime="text/html"
                )
                
                # Show HTML preview
                with st.expander("HTML Preview"):
                    st.components.v1.html(html_content, height=500)
            
            elif output_format == "PDF":
                pdf_data = convert_to_pdf(html_content)
                if pdf_data:
                    st.download_button(
                        label="Download PDF",
                        data=pdf_data,
                        file_name=f"{uploaded_file.name.split('.')[0]}.pdf",
                        mime="application/pdf"
                    )
            
            elif output_format == "DOCX":
                docx_data = convert_to_docx_with_style(html_content)
                st.download_button(
                    label="Download DOCX",
                    data=docx_data,
                    file_name=f"{uploaded_file.name.split('.')[0]}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

if __name__ == "__main__":
    main()
